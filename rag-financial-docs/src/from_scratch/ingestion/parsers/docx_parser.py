"""
Parse DOCX files using python-docx, preserving document structure.

DOCX is easier than PDF because the format is well-defined XML under the hood.
python-docx gives us clean access to paragraphs, headings, and tables.

Key decisions:
- Headings are preserved as markdown-style markers (# Heading, ## Heading)
  so chunkers can use them as natural split points.
- Tables are extracted separately, same interface as pdf_parser.
- Empty paragraphs are dropped to reduce noise.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

import docx
from docx.oxml.ns import qn


# Heading style names in python-docx map to these levels
_HEADING_STYLES = {
    "Heading 1": "#",
    "Heading 2": "##",
    "Heading 3": "###",
    "Heading 4": "####",
    "Title": "#",
}


@dataclass
class ParsedDocxDocument:
    """Result of parsing one DOCX file."""

    text: str                                      # Full text with heading markers
    tables: list[list[list[str]]] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    source_file: str = ""

    def tables_as_text(self) -> str:
        """Convert all tables to readable text."""
        if not self.tables:
            return ""
        parts = []
        for i, table in enumerate(self.tables):
            rows = [" | ".join(cell or "" for cell in row) for row in table]
            parts.append(f"[Table {i + 1}]\n" + "\n".join(rows))
        return "\n\n".join(parts)


def _paragraph_to_text(para: docx.text.paragraph.Paragraph) -> str | None:
    """
    Convert a paragraph to a string, applying heading markers.
    Returns None for empty paragraphs (they are skipped).
    """
    text = para.text.strip()
    if not text:
        return None

    style_name = para.style.name if para.style else ""
    prefix = _HEADING_STYLES.get(style_name, "")

    if prefix:
        return f"{prefix} {text}"
    return text


def _extract_table(table: docx.table.Table) -> list[list[str]]:
    """Extract a docx table into a list of rows (each row is a list of cell strings)."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)
    return rows


def parse_docx(
    path: str | Path,
    metadata: dict | None = None,
) -> ParsedDocxDocument:
    """
    Parse a DOCX file into text and tables.

    Args:
        path: Path to the .docx file.
        metadata: Caller-supplied metadata dict.

    Returns:
        ParsedDocxDocument with full text (headings as markdown) and tables.

    Example:
        doc = parse_docx(
            "data/raw/policies/risk_management.docx",
            metadata={
                "doc_type": "policy",
                "confidentiality": "internal",
            },
        )
        print(doc.text[:500])
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    metadata = dict(metadata or {})
    metadata["source_file"] = str(path)
    metadata["filename"] = path.name

    document = docx.Document(str(path))

    # --- Extract paragraphs ---
    text_parts: list[str] = []
    for para in document.paragraphs:
        line = _paragraph_to_text(para)
        if line:
            text_parts.append(line)

    full_text = "\n\n".join(text_parts)

    # --- Extract tables ---
    tables = [_extract_table(tbl) for tbl in document.tables]

    return ParsedDocxDocument(
        text=full_text,
        tables=tables,
        metadata=metadata,
        source_file=str(path),
    )
