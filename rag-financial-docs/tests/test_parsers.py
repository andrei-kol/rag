"""
Tests for PDF, DOCX, and table parsers.

We test the parsers without any real documents by creating minimal fixtures
on the fly (using reportlab for PDF, python-docx for DOCX).
This keeps tests fast and self-contained — no large files in the repo.
"""

from __future__ import annotations

import io
from pathlib import Path

import pytest

# ---------------------------------------------------------------------------
# Fixtures — create minimal test documents in memory / tmp
# ---------------------------------------------------------------------------


@pytest.fixture()
def simple_pdf(tmp_path: Path) -> Path:
    """Create a minimal two-page PDF with one table."""
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table
    from reportlab.lib.styles import getSampleStyleSheet

    styles = getSampleStyleSheet()
    out = tmp_path / "test.pdf"

    doc = SimpleDocTemplate(str(out), pagesize=A4)
    story = [
        Paragraph("Acme Corp Annual Report 2023", styles["Title"]),
        Paragraph("Revenue for the year was $42.5 million, up 18% year-over-year.", styles["Normal"]),
        Spacer(1, 20),
        Table(
            [
                ["Metric", "Q1", "Q2", "Q3", "Q4"],
                ["Revenue ($M)", "9.5", "10.1", "11.2", "11.7"],
                ["Net Income ($M)", "1.1", "1.4", "1.7", "1.9"],
            ]
        ),
        Paragraph("This report is confidential.", styles["Normal"]),
    ]
    doc.build(story)
    return out


@pytest.fixture()
def simple_docx(tmp_path: Path) -> Path:
    """Create a minimal DOCX with headings and a table."""
    import docx as python_docx

    out = tmp_path / "test.docx"
    document = python_docx.Document()
    document.add_heading("Risk Management Policy", level=1)
    document.add_heading("1. Purpose", level=2)
    document.add_paragraph("This policy defines risk management procedures.")
    document.add_heading("2. Categories", level=2)
    document.add_paragraph("Market, credit, and operational risk.")

    tbl = document.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "Risk"
    tbl.rows[0].cells[1].text = "Action"
    tbl.rows[1].cells[0].text = "High"
    tbl.rows[1].cells[1].text = "Escalate"

    document.save(str(out))
    return out


# ---------------------------------------------------------------------------
# PDF parser tests
# ---------------------------------------------------------------------------


class TestPdfParser:
    def test_returns_parsed_document(self, simple_pdf):
        from src.from_scratch.ingestion.parsers.pdf_parser import parse_pdf, ParsedDocument

        result = parse_pdf(simple_pdf)
        assert isinstance(result, ParsedDocument)

    def test_extracts_text(self, simple_pdf):
        from src.from_scratch.ingestion.parsers.pdf_parser import parse_pdf

        result = parse_pdf(simple_pdf)
        assert "Acme Corp" in result.full_text
        assert "Revenue" in result.full_text

    def test_has_pages(self, simple_pdf):
        from src.from_scratch.ingestion.parsers.pdf_parser import parse_pdf

        result = parse_pdf(simple_pdf)
        assert result.num_pages >= 1

    def test_page_markers_in_full_text(self, simple_pdf):
        from src.from_scratch.ingestion.parsers.pdf_parser import parse_pdf

        result = parse_pdf(simple_pdf)
        assert "[Page 1]" in result.full_text

    def test_metadata_merged(self, simple_pdf):
        from src.from_scratch.ingestion.parsers.pdf_parser import parse_pdf

        meta = {"client_id": "acme_corp", "doc_type": "financial_report", "year": 2023}
        result = parse_pdf(simple_pdf, metadata=meta)

        assert result.metadata["client_id"] == "acme_corp"
        assert result.metadata["doc_type"] == "financial_report"
        assert result.metadata["year"] == 2023
        # Parser adds these automatically
        assert "source_file" in result.metadata
        assert "filename" in result.metadata
        assert result.metadata["page_count"] == result.num_pages

    def test_tables_extracted(self, simple_pdf):
        from src.from_scratch.ingestion.parsers.pdf_parser import parse_pdf

        result = parse_pdf(simple_pdf, extract_tables=True)
        # Our fixture has one table — pdfplumber should find it
        assert len(result.all_tables) >= 1

    def test_no_tables_when_disabled(self, simple_pdf):
        from src.from_scratch.ingestion.parsers.pdf_parser import parse_pdf

        result = parse_pdf(simple_pdf, extract_tables=False)
        assert result.all_tables == []

    def test_file_not_found_raises(self, tmp_path):
        from src.from_scratch.ingestion.parsers.pdf_parser import parse_pdf

        with pytest.raises(FileNotFoundError):
            parse_pdf(tmp_path / "nonexistent.pdf")

    def test_tables_as_text(self, simple_pdf):
        from src.from_scratch.ingestion.parsers.pdf_parser import parse_pdf

        result = parse_pdf(simple_pdf, extract_tables=True)
        if result.all_tables:
            text = result.tables_as_text()
            assert "[Table 1]" in text
            assert "|" in text  # pipe-delimited


# ---------------------------------------------------------------------------
# DOCX parser tests
# ---------------------------------------------------------------------------


class TestDocxParser:
    def test_returns_parsed_docx_document(self, simple_docx):
        from src.from_scratch.ingestion.parsers.docx_parser import parse_docx, ParsedDocxDocument

        result = parse_docx(simple_docx)
        assert isinstance(result, ParsedDocxDocument)

    def test_headings_have_markdown_markers(self, simple_docx):
        from src.from_scratch.ingestion.parsers.docx_parser import parse_docx

        result = parse_docx(simple_docx)
        assert "# Risk Management Policy" in result.text
        assert "## 1. Purpose" in result.text

    def test_body_text_preserved(self, simple_docx):
        from src.from_scratch.ingestion.parsers.docx_parser import parse_docx

        result = parse_docx(simple_docx)
        assert "This policy defines risk management procedures." in result.text

    def test_tables_extracted(self, simple_docx):
        from src.from_scratch.ingestion.parsers.docx_parser import parse_docx

        result = parse_docx(simple_docx)
        assert len(result.tables) == 1
        assert result.tables[0][0] == ["Risk", "Action"]
        assert result.tables[0][1] == ["High", "Escalate"]

    def test_metadata_merged(self, simple_docx):
        from src.from_scratch.ingestion.parsers.docx_parser import parse_docx

        meta = {"doc_type": "policy", "confidentiality": "internal"}
        result = parse_docx(simple_docx, metadata=meta)

        assert result.metadata["doc_type"] == "policy"
        assert result.metadata["confidentiality"] == "internal"
        assert "source_file" in result.metadata

    def test_file_not_found_raises(self, tmp_path):
        from src.from_scratch.ingestion.parsers.docx_parser import parse_docx

        with pytest.raises(FileNotFoundError):
            parse_docx(tmp_path / "nonexistent.docx")


# ---------------------------------------------------------------------------
# Table parser tests
# ---------------------------------------------------------------------------


class TestTableParser:
    RAW_TABLE = [
        ["Metric", "Q1", "Q2"],
        ["Revenue", "9.5", "10.1"],
        [None, "9.5", None],   # None cells — common in real PDFs
    ]

    def test_clean_table_removes_none(self):
        from src.from_scratch.ingestion.parsers.table_parser import clean_table

        cleaned = clean_table(self.RAW_TABLE)
        for row in cleaned:
            for cell in row:
                assert cell is not None
                assert isinstance(cell, str)

    def test_clean_table_strips_whitespace(self):
        from src.from_scratch.ingestion.parsers.table_parser import clean_table

        raw = [["  Revenue  ", " 9.5 "]]
        cleaned = clean_table(raw)
        assert cleaned[0][0] == "Revenue"
        assert cleaned[0][1] == "9.5"

    def test_table_to_markdown_has_header_separator(self):
        from src.from_scratch.ingestion.parsers.table_parser import table_to_markdown, clean_table

        cleaned = clean_table(self.RAW_TABLE)
        md = table_to_markdown(cleaned)
        lines = md.strip().split("\n")
        # Line 0: header, Line 1: separator (---), Line 2+: data rows
        assert "---" in lines[1]

    def test_table_to_markdown_with_title(self):
        from src.from_scratch.ingestion.parsers.table_parser import table_to_markdown, clean_table

        cleaned = clean_table(self.RAW_TABLE)
        md = table_to_markdown(cleaned, title="Revenue Breakdown")
        assert "[Table: Revenue Breakdown]" in md

    def test_table_to_markdown_empty_returns_empty(self):
        from src.from_scratch.ingestion.parsers.table_parser import table_to_markdown

        assert table_to_markdown([]) == ""

    def test_make_table_chunk_sets_content_type(self):
        from src.from_scratch.ingestion.parsers.table_parser import make_table_chunk

        chunk = make_table_chunk(
            raw_table=self.RAW_TABLE,
            metadata={"client_id": "acme_corp"},
            title="Revenue",
        )
        assert chunk.metadata["content_type"] == "table"
        assert chunk.metadata["table_title"] == "Revenue"
        assert chunk.metadata["client_id"] == "acme_corp"

    def test_make_table_chunk_does_not_mutate_input_metadata(self):
        from src.from_scratch.ingestion.parsers.table_parser import make_table_chunk

        original = {"client_id": "acme_corp"}
        make_table_chunk(raw_table=self.RAW_TABLE, metadata=original)
        # original should not have content_type added to it
        assert "content_type" not in original

    def test_make_table_chunk_text_contains_pipe(self):
        from src.from_scratch.ingestion.parsers.table_parser import make_table_chunk

        chunk = make_table_chunk(raw_table=self.RAW_TABLE)
        assert "|" in chunk.text_representation
